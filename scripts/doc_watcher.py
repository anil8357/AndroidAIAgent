"""
Document Watcher — Auto-converts documents to agent-readable Markdown.

Drop any PDF, DOCX, or image file into docs/input/ and this script will:
1. Detect the new file
2. Convert it to Markdown
3. Save the result in docs/parsed/<filename>.md
4. The OpenCode agent can then read docs/parsed/ like any other text file

Supported formats:
- PDF  → pymupdf4llm (high quality LLM-friendly markdown)
- DOCX → python-docx (preserves headings, lists, tables)
- Images (PNG, JPG, JPEG, BMP, TIFF) → Gemini Vision API via LiteLLM

Usage:
    python scripts/doc_watcher.py          # Watch mode (runs continuously)
    python scripts/doc_watcher.py --once   # Process all files once and exit

Requirements:
    pip install pymupdf4llm python-docx watchdog requests Pillow
"""

import os
import sys
import time
import json
import base64
import argparse
from pathlib import Path

# --- Configuration ---
PROJECT_ROOT = Path(__file__).resolve().parent.parent
INPUT_DIR = PROJECT_ROOT / "docs" / "input"
OUTPUT_DIR = PROJECT_ROOT / "docs" / "parsed"
LITELLM_BASE_URL = os.environ.get("LITELLM_BASE_URL", "http://localhost:4000/v1")
LITELLM_API_KEY = os.environ.get("LITELLM_API_KEY", "sk-1234")
VISION_MODEL = os.environ.get("VISION_MODEL", "coder")  # Gemini 2.5 Flash via LiteLLM

# --- Vision request tuning ---
# Truncation ("stops reading in the middle") happens when the model hits max_tokens.
# We request a high cap AND auto-continue when the response is cut off.
VISION_MAX_TOKENS = int(os.environ.get("VISION_MAX_TOKENS", "8192"))
# Big images on slow fallback models can take well over a minute.
VISION_TIMEOUT = int(os.environ.get("VISION_TIMEOUT", "180"))
# Transient 429/5xx/connection errors are retried with exponential backoff.
VISION_MAX_RETRIES = int(os.environ.get("VISION_MAX_RETRIES", "3"))
# If the model output is cut off (finish_reason == "length"), keep asking for more.
VISION_MAX_CONTINUATIONS = int(os.environ.get("VISION_MAX_CONTINUATIONS", "4"))

# Supported file extensions
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx", ".doc"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}
ALL_SUPPORTED = PDF_EXTENSIONS | DOCX_EXTENSIONS | IMAGE_EXTENSIONS


def ensure_dirs():
    """Create input/output directories if they don't exist."""
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _wait_until_stable(path: Path, checks: int = 3, interval: float = 0.7) -> None:
    """Block until a file's size stops changing, so we don't process a partial write."""
    last_size = -1
    stable = 0
    for _ in range(20):  # cap total wait at ~14s
        try:
            size = path.stat().st_size
        except OSError:
            time.sleep(interval)
            continue
        if size == last_size and size > 0:
            stable += 1
            if stable >= checks:
                return
        else:
            stable = 0
        last_size = size
        time.sleep(interval)


def convert_pdf_to_md(input_path: Path) -> str:
    """Convert PDF to Markdown using pymupdf4llm."""
    try:
        import pymupdf4llm
        md_text = pymupdf4llm.to_markdown(str(input_path))
        return md_text
    except ImportError:
        # Fallback to basic pymupdf text extraction
        import fitz  # PyMuPDF
        doc = fitz.open(str(input_path))
        text_parts = []
        for page_num, page in enumerate(doc, 1):
            text_parts.append(f"## Page {page_num}\n\n{page.get_text()}")
        doc.close()
        return "\n\n---\n\n".join(text_parts)


def convert_docx_to_md(input_path: Path) -> str:
    """Convert DOCX to Markdown using python-docx."""
    from docx import Document
    from docx.table import Table

    doc = Document(str(input_path))
    lines = []
    
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('}p'):
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if not text:
                        lines.append("")
                        break
                    
                    style = para.style.name if para.style else ""
                    
                    if "Heading 1" in style:
                        lines.append(f"# {text}")
                    elif "Heading 2" in style:
                        lines.append(f"## {text}")
                    elif "Heading 3" in style:
                        lines.append(f"### {text}")
                    elif "Heading 4" in style:
                        lines.append(f"#### {text}")
                    elif "List" in style or "Bullet" in style:
                        lines.append(f"- {text}")
                    elif "Number" in style:
                        lines.append(f"1. {text}")
                    else:
                        # Check for bold/italic
                        if para.runs and all(r.bold for r in para.runs if r.text.strip()):
                            lines.append(f"**{text}**")
                        else:
                            lines.append(text)
                    lines.append("")
                    break
        
        # Handle tables
        elif element.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._element == element:
                    lines.append(convert_table_to_md(table))
                    lines.append("")
                    break
    
    return "\n".join(lines)


def convert_table_to_md(table) -> str:
    """Convert a docx table to markdown table format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    
    if len(rows) >= 1:
        # Add header separator after first row
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)
    
    return "\n".join(rows)


MIME_MAP = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
    ".webp": "image/webp",
}


class VisionError(Exception):
    """Raised when the vision model cannot produce usable output."""


def encode_image(input_path: Path) -> tuple[str, str, int, int]:
    """Return (base64_data, mime_type, width, height) for an image file."""
    from PIL import Image

    with open(input_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    mime_type = MIME_MAP.get(input_path.suffix.lower(), "image/png")

    with Image.open(input_path) as img:
        width, height = img.size

    return image_data, mime_type, width, height


def call_vision_model(prompt_text: str, image_data: str, mime_type: str) -> str:
    """
    Call the vision model via LiteLLM with full reliability handling:

    - Retries transient failures (429 / 5xx / connection / timeout) with backoff.
    - Auto-continues when the response is cut off at max_tokens (finish_reason == "length"),
      so long documents/screenshots are no longer truncated mid-output.
    - Detects when a non-vision fallback model rejected the image and raises a clear error.

    Returns the full concatenated text content. Raises VisionError on unrecoverable failure.
    """
    import requests

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LITELLM_API_KEY}",
    }

    base_messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt_text},
                {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{image_data}"}},
            ],
        }
    ]

    full_text = ""
    used_model = None
    continuations = 0

    while True:
        payload = {
            "model": VISION_MODEL,
            "messages": base_messages,
            "max_tokens": VISION_MAX_TOKENS,
            "temperature": 0.1,
        }

        # Ask the model to resume exactly where it left off on a truncated response.
        if full_text:
            payload["messages"] = base_messages + [
                {"role": "assistant", "content": full_text},
                {
                    "role": "user",
                    "content": "Continue exactly where you stopped. Do not repeat earlier content.",
                },
            ]

        result = _post_with_retries(requests, headers, payload, input_label="image")

        choice = result["choices"][0]
        used_model = result.get("model", used_model)
        chunk = (choice.get("message") or {}).get("content") or ""
        full_text += chunk

        finish_reason = choice.get("finish_reason")
        if finish_reason == "length" and continuations < VISION_MAX_CONTINUATIONS:
            continuations += 1
            print(f"  ↪️  Output truncated — requesting continuation {continuations}/{VISION_MAX_CONTINUATIONS}")
            continue
        break

    if not full_text.strip():
        raise VisionError(
            "Vision model returned empty content. The active model "
            f"('{used_model or VISION_MODEL}') may not support image input "
            "(Gemini's daily free quota may be exhausted, causing fallback to a text-only model)."
        )

    return full_text


def _post_with_retries(requests_mod, headers, payload, input_label: str) -> dict:
    """POST to LiteLLM with exponential backoff for transient errors."""
    last_err = None
    for attempt in range(1, VISION_MAX_RETRIES + 1):
        try:
            response = requests_mod.post(
                f"{LITELLM_BASE_URL}/chat/completions",
                headers=headers,
                json=payload,
                timeout=VISION_TIMEOUT,
            )
            if response.status_code in (429, 500, 502, 503, 504):
                last_err = VisionError(f"HTTP {response.status_code}: {response.text[:300]}")
                raise last_err
            response.raise_for_status()
            return response.json()
        except requests_mod.exceptions.ConnectionError as e:
            raise VisionError(
                f"Could not connect to LiteLLM at {LITELLM_BASE_URL}. "
                "Start the proxy first: `docker compose up -d` "
                "(or `litellm --config litellm_config.yaml --port 4000`)."
            ) from e
        except requests_mod.exceptions.HTTPError as e:
            # 4xx other than 429 — usually the model rejected the image (non-vision fallback).
            status = getattr(e.response, "status_code", "?")
            body = getattr(e.response, "text", "")[:300]
            raise VisionError(
                f"Vision request rejected (HTTP {status}). The active model may not accept "
                f"images. Details: {body}"
            ) from e
        except (requests_mod.exceptions.Timeout, VisionError) as e:
            last_err = e
            if attempt < VISION_MAX_RETRIES:
                backoff = 2 ** attempt
                print(f"  ⏳ {input_label} attempt {attempt} failed ({e}); retrying in {backoff}s")
                time.sleep(backoff)
            continue

    raise VisionError(f"Vision request failed after {VISION_MAX_RETRIES} attempts: {last_err}")


IMAGE_EXTRACTION_PROMPT = (
    "You are a document analyst. Extract ALL information from this image "
    "and convert it to well-structured Markdown. Include:\n"
    "- All text content (preserve headings, lists, tables)\n"
    "- Description of any diagrams, flowcharts, or wireframes\n"
    "- Any data in tables or charts\n"
    "- UI mockup descriptions (if applicable)\n"
    "Be thorough and complete — do not stop early. This will be used as a reference document."
)


def convert_image_to_md(input_path: Path) -> str:
    """Image conversion is no longer supported by @doc-reader.

    Images are skipped. Users should describe their UI needs to @ui-builder instead.
    This function remains only as a stub in case screenshot_to_ui.py's imports reference it.
    """
    return (
        f"# Image: {input_path.name}\n\n"
        f"> ⚠️ Image parsing is not supported. Describe your UI to @ui-builder instead.\n"
    )


def process_file(input_path: Path) -> bool:
    """Process a single file and save the markdown output."""
    ext = input_path.suffix.lower()
    output_path = OUTPUT_DIR / f"{input_path.stem}.md"
    
    if ext not in ALL_SUPPORTED:
        print(f"  ⏭️  Skipped (unsupported): {input_path.name}")
        return False
    
    print(f"  📄 Processing: {input_path.name} → {output_path.name}")
    
    try:
        if ext in PDF_EXTENSIONS:
            md_content = convert_pdf_to_md(input_path)
        elif ext in DOCX_EXTENSIONS:
            md_content = convert_docx_to_md(input_path)
        elif ext in IMAGE_EXTENSIONS:
            print(f"  ⏭️  Skipped (images not supported — describe to @ui-builder): {input_path.name}")
            return False
        else:
            return False
        
        # Add metadata header
        metadata = (
            f"<!-- \n"
            f"  Source: {input_path.name}\n"
            f"  Converted: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"  Type: {ext}\n"
            f"-->\n\n"
        )
        
        output_path.write_text(metadata + md_content, encoding="utf-8")
        print(f"  ✅ Saved: {output_path.name} ({output_path.stat().st_size / 1024:.1f} KB)")
        return True
        
    except Exception as e:
        print(f"  ❌ Error processing {input_path.name}: {e}")
        return False


def process_all_existing():
    """Process all files currently in the input directory."""
    files = [f for f in INPUT_DIR.iterdir() if f.is_file() and f.suffix.lower() in ALL_SUPPORTED]
    
    if not files:
        print(f"📂 No supported files found in {INPUT_DIR}")
        print(f"   Supported: {', '.join(sorted(PDF_EXTENSIONS | DOCX_EXTENSIONS))}")
        print(f"   (Images are not parsed — describe your UI to @ui-builder instead)")
        return
    
    print(f"📂 Found {len(files)} file(s) to process:\n")
    
    success_count = 0
    for file_path in sorted(files):
        if process_file(file_path):
            success_count += 1
    
    print(f"\n✅ Done: {success_count}/{len(files)} files converted to Markdown")
    print(f"📁 Output: {OUTPUT_DIR}")


def watch_directory():
    """Watch the input directory for new files and process them automatically."""
    try:
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler
    except ImportError:
        print("❌ watchdog not installed. Install with: pip install watchdog")
        print("   Falling back to polling mode...\n")
        watch_directory_polling()
        return
    
    class DocHandler(FileSystemEventHandler):
        def on_created(self, event):
            if event.is_directory:
                return
            file_path = Path(event.src_path)
            if file_path.suffix.lower() in ALL_SUPPORTED:
                # Wait for the file to be fully written (large PDFs/images take time).
                _wait_until_stable(file_path)
                print(f"\n🔔 New file detected!")
                process_file(file_path)
    
    observer = Observer()
    observer.schedule(DocHandler(), str(INPUT_DIR), recursive=False)
    observer.start()
    
    print(f"👁️  Watching: {INPUT_DIR}")
    print(f"📁 Output:   {OUTPUT_DIR}")
    print(f"📎 Drop PDF, DOCX, or images into the input folder.")
    print(f"   Press Ctrl+C to stop.\n")
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
        print("\n🛑 Watcher stopped.")
    observer.join()


def watch_directory_polling():
    """Fallback: poll-based watcher if watchdog is not available."""
    seen_files = set()
    
    # Mark existing files as already seen
    for f in INPUT_DIR.iterdir():
        if f.is_file():
            seen_files.add(f.name)
    
    print(f"👁️  Watching (polling): {INPUT_DIR}")
    print(f"📁 Output: {OUTPUT_DIR}")
    print(f"📎 Drop PDF, DOCX, or images into the input folder.")
    print(f"   Press Ctrl+C to stop.\n")
    
    try:
        while True:
            for f in INPUT_DIR.iterdir():
                if f.is_file() and f.name not in seen_files:
                    seen_files.add(f.name)
                    if f.suffix.lower() in ALL_SUPPORTED:
                        _wait_until_stable(f)  # Wait for write to complete
                        print(f"\n🔔 New file detected!")
                        process_file(f)
            time.sleep(2)
    except KeyboardInterrupt:
        print("\n🛑 Watcher stopped.")


def main():
    parser = argparse.ArgumentParser(description="Convert documents to agent-readable Markdown")
    parser.add_argument("--once", action="store_true", help="Process existing files and exit")
    parser.add_argument("--file", type=str, help="Process a single file")
    args = parser.parse_args()
    
    ensure_dirs()
    
    if args.file:
        file_path = Path(args.file)
        if not file_path.exists():
            print(f"❌ File not found: {file_path}")
            sys.exit(1)
        process_file(file_path)
    elif args.once:
        process_all_existing()
    else:
        # Process existing files first, then watch
        process_all_existing()
        print("\n" + "=" * 50)
        print("Switching to watch mode...\n")
        watch_directory()


if __name__ == "__main__":
    main()
